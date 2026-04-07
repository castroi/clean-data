import pytest
from pathlib import Path

import fitz
from docx import Document

from processor.pipeline import CleaningPipeline


@pytest.fixture(scope="module")
def pipeline():
    return CleaningPipeline()


def _create_test_pdf(path: Path, text: str) -> None:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), text)
    doc.set_metadata({"author": "John Smith"})
    doc.save(str(path))
    doc.close()


def _create_test_docx(path: Path, paragraphs: list[str]) -> None:
    doc = Document()
    doc.core_properties.author = "John Smith"
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(str(path))


def test_pipeline_pdf(pipeline: CleaningPipeline, tmp_path: Path):
    input_path = tmp_path / "test.pdf"
    _create_test_pdf(input_path, "Contact John Smith at john@example.com")

    cleaned = pipeline.process(input_path)

    assert cleaned.exists()
    doc = fitz.open(str(cleaned))
    text = doc[0].get_text()
    doc.close()
    assert "John" not in text
    assert "john@example.com" not in text


def test_pipeline_docx(pipeline: CleaningPipeline, tmp_path: Path):
    input_path = tmp_path / "test.docx"
    _create_test_docx(input_path, ["Contact John Smith", "Email: john@example.com"])

    cleaned = pipeline.process(input_path)

    assert cleaned.exists()
    doc = Document(str(cleaned))
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "John" not in full_text
    assert "john@example.com" not in full_text


def test_pipeline_unsupported_type(pipeline: CleaningPipeline, tmp_path: Path):
    input_path = tmp_path / "file.xlsx"
    input_path.write_text("dummy")

    with pytest.raises(ValueError, match="Unsupported"):
        pipeline.process(input_path)


def test_pipeline_cleanup(pipeline: CleaningPipeline, tmp_path: Path):
    input_path = tmp_path / "test.pdf"
    _create_test_pdf(input_path, "Email: test@example.com")

    cleaned = pipeline.process(input_path)
    assert cleaned.exists()

    pipeline.cleanup([input_path, cleaned])
    assert not input_path.exists()
    assert not cleaned.exists()


def test_pipeline_output_named_correctly(pipeline: CleaningPipeline, tmp_path: Path):
    input_path = tmp_path / "report.pdf"
    _create_test_pdf(input_path, "Hello world")

    cleaned = pipeline.process(input_path)
    assert cleaned.name == "cleaned_report.pdf"
