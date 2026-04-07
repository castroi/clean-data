import pytest
from pathlib import Path

from docx import Document

from processor.pii_detector import PIIDetector
from processor.docx_cleaner import DOCXCleaner


@pytest.fixture(scope="module")
def detector():
    return PIIDetector()


@pytest.fixture(scope="module")
def cleaner(detector):
    return DOCXCleaner(detector)


def _create_test_docx(path: Path, paragraphs: list[str], author: str = "John Smith") -> None:
    doc = Document()
    doc.core_properties.author = author
    doc.core_properties.title = "Secret Report"
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(str(path))


def _create_test_docx_with_table(path: Path, rows: list[list[str]]) -> None:
    doc = Document()
    doc.core_properties.author = "John Smith"
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = cell_text
    doc.save(str(path))


def test_removes_pii_from_paragraphs(cleaner: DOCXCleaner, tmp_path: Path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    _create_test_docx(input_path, ["Contact John Smith", "Email: john@example.com"])

    stats = cleaner.clean(input_path, output_path)

    doc = Document(str(output_path))
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "John" not in full_text
    assert "john@example.com" not in full_text
    assert stats["pii_items_removed"] > 0


def test_removes_pii_from_tables(cleaner: DOCXCleaner, tmp_path: Path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    _create_test_docx_with_table(input_path, [
        ["Name", "Email"],
        ["John Smith", "john@example.com"],
    ])

    cleaner.clean(input_path, output_path)

    doc = Document(str(output_path))
    all_text = ""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += cell.text + " "
    assert "John" not in all_text
    assert "john@example.com" not in all_text


def test_strips_metadata(cleaner: DOCXCleaner, tmp_path: Path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    _create_test_docx(input_path, ["Hello world"])

    cleaner.clean(input_path, output_path)

    doc = Document(str(output_path))
    assert doc.core_properties.author in (None, "")
    assert doc.core_properties.title in (None, "")


def test_returns_stats(cleaner: DOCXCleaner, tmp_path: Path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    _create_test_docx(input_path, ["Email: test@example.com"])

    stats = cleaner.clean(input_path, output_path)

    assert "sections_processed" in stats
    assert "pii_items_removed" in stats
    assert "metadata_stripped" in stats
    assert stats["metadata_stripped"] is True


def test_output_is_valid_docx(cleaner: DOCXCleaner, tmp_path: Path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    _create_test_docx(input_path, ["Contact John Smith"])

    cleaner.clean(input_path, output_path)

    assert output_path.exists()
    doc = Document(str(output_path))
    assert len(doc.paragraphs) > 0


def test_preserves_non_pii_text(cleaner: DOCXCleaner, tmp_path: Path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    _create_test_docx(input_path, ["The weather is nice today"])

    cleaner.clean(input_path, output_path)

    doc = Document(str(output_path))
    text = doc.paragraphs[0].text
    assert "weather" in text
    assert "nice" in text


def test_removes_israeli_id_from_docx(cleaner: DOCXCleaner, tmp_path: Path):
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    _create_test_docx(input_path, ["ID number: 123456782"])

    cleaner.clean(input_path, output_path)

    doc = Document(str(output_path))
    text = doc.paragraphs[0].text
    assert "123456782" not in text
