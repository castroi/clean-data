import logging
from pathlib import Path

from docx import Document

from processor.metadata import strip_docx_metadata
from processor.pii_detector import PIIDetector

logger = logging.getLogger(__name__)


class DOCXCleaner:
    """Cleans PII from DOCX documents and strips metadata."""

    def __init__(self, pii_detector: PIIDetector) -> None:
        self._detector = pii_detector

    def _clean_text_in_paragraph(self, paragraph) -> int:
        """Clean PII from a paragraph's runs. Returns count of PII items found."""
        full_text = paragraph.text
        if not full_text.strip():
            return 0

        entities = self._detector.detect_entities(full_text)
        if not entities:
            return 0

        cleaned = self._detector.detect_and_remove(full_text)

        # Preserve formatting by putting cleaned text in first run
        # and clearing subsequent runs
        if paragraph.runs:
            paragraph.runs[0].text = cleaned
            for run in paragraph.runs[1:]:
                run.text = ""
        elif full_text != cleaned:
            # Paragraph has text but no runs (rare) — set directly
            paragraph.text = cleaned

        return len(entities)

    def clean(self, input_path: Path, output_path: Path) -> dict:
        """Clean PII from a DOCX file and save the result.

        Returns stats dict with sections_processed, pii_items_removed,
        and metadata_stripped.
        """
        doc = Document(str(input_path))
        total_pii_removed = 0
        sections_processed = 0

        # Clean paragraphs
        for paragraph in doc.paragraphs:
            total_pii_removed += self._clean_text_in_paragraph(paragraph)
            sections_processed += 1

        # Clean tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        total_pii_removed += self._clean_text_in_paragraph(paragraph)
                        sections_processed += 1

        # Clean headers and footers
        for section in doc.sections:
            for header_footer in (section.header, section.footer):
                if header_footer is not None:
                    for paragraph in header_footer.paragraphs:
                        total_pii_removed += self._clean_text_in_paragraph(paragraph)
                        sections_processed += 1

        # Strip metadata
        strip_docx_metadata(doc)

        doc.save(str(output_path))

        logger.info(
            "DOCX cleaned: %d sections, %d PII items removed",
            sections_processed,
            total_pii_removed,
        )

        return {
            "sections_processed": sections_processed,
            "pii_items_removed": total_pii_removed,
            "metadata_stripped": True,
        }
